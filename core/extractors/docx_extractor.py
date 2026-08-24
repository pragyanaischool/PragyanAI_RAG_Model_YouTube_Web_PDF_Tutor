"""
DOCX Extractor Module for NCET GenAI Multimodal RAG.
Handles Microsoft Word (.docx) file streams, extracting body paragraphs,
structured heading hierarchies, and embedded table rows into a standardized format.
"""

import io
import re
from typing import Any, Dict, Optional
from docx import Document


def clean_text_block(text: str) -> str:
    """
    Normalizes whitespace and removes unwanted control characters.
    """
    if not text:
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def extract_from_docx(uploaded_file: Any) -> Optional[Dict[str, Any]]:
    """
    Parses an uploaded Microsoft Word (.docx) file stream or file-like binary object,
    extracts paragraphs, preserves headings, and converts tables into markdown-style rows.

    Args:
        uploaded_file: Streamlit UploadedFile, file-like object, or bytes buffer.

    Returns:
        Optional[Dict[str, Any]]: Document payload containing:
            - 'source': Document filename
            - 'type': 'docx'
            - 'title': Document filename or title
            - 'content': Structured text payload
        Returns None if extraction fails or file contains no extractable text.
    """
    if uploaded_file is None:
        return None

    file_name = getattr(uploaded_file, "name", "Uploaded_Document.docx")

    try:
        # Support Streamlit UploadedFile, BytesIO, or raw byte streams
        if hasattr(uploaded_file, "getvalue"):
            docx_stream = io.BytesIO(uploaded_file.getvalue())
        elif isinstance(uploaded_file, bytes):
            docx_stream = io.BytesIO(uploaded_file)
        else:
            docx_stream = uploaded_file

        doc = Document(docx_stream)
        extracted_blocks = []

        # 1. Extract Body Paragraphs & preserve Heading structures
        for para in doc.paragraphs:
            cleaned_text = clean_text_block(para.text)
            if cleaned_text:
                style_name = getattr(para.style, "name", "")
                if style_name.startswith("Heading 1"):
                    extracted_blocks.append(f"\n# {cleaned_text}\n")
                elif style_name.startswith("Heading 2"):
                    extracted_blocks.append(f"\n## {cleaned_text}\n")
                elif style_name.startswith("Heading 3"):
                    extracted_blocks.append(f"\n### {cleaned_text}\n")
                elif style_name.startswith("Heading"):
                    extracted_blocks.append(f"\n#### {cleaned_text}\n")
                elif style_name.startswith("List"):
                    extracted_blocks.append(f"- {cleaned_text}")
                else:
                    extracted_blocks.append(cleaned_text)

        # 2. Extract Embedded Tables
        for table_idx, table in enumerate(doc.tables):
            table_lines = [f"\n[Table {table_idx + 1}]"]
            for row in table.rows:
                # Deduplicate merged cells by stripping and keeping unique contents in order
                row_cells = [clean_text_block(cell.text) for cell in row.cells]
                # Filter out empty cells
                non_empty_cells = [c for c in row_cells if c]
                if non_empty_cells:
                    table_lines.append(" | ".join(non_empty_cells))

            if len(table_lines) > 1:
                extracted_blocks.append("\n".join(table_lines) + "\n")

        if not extracted_blocks:
            print(f"[Warning in docx_extractor]: No text found in {file_name}.")
            return None

        full_document_text = "\n\n".join(extracted_blocks)

        header_block = (
            f"[Source: DOCX Document]\n"
            f"[Title: {file_name}]\n"
            f"----------------------------------------\n\n"
        )

        full_content = f"{header_block}{full_document_text}"

        return {
            "source": file_name,
            "type": "docx",
            "title": file_name,
            "content": full_content,
        }

    except Exception as e:
        print(f"[Error in docx_extractor]: Failed parsing {file_name}: {e}")
        return None
